def export_pdf(fig):
    path = "exports/chart.pdf"
    fig.write_image(path)
    return path

def export_jpeg(fig):
    path = "exports/chart.jpeg"
    fig.write_image(path)
    return path

def export_word(fig):
    img_path = "exports/chart.jpeg"
    fig.write_image(img_path)

    from docx import Document
    doc = Document()
    doc.add_heading("Datum Visualization Report", level=1)
    doc.add_picture(img_path)

    word_path = "exports/chart.docx"
    doc.save(word_path)
    return word_path
